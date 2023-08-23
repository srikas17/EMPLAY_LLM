#!/usr/bin/env python
# coding: utf-8

# In[1]:


get_ipython().system('pip install python-docx translate streamlit')


# In[2]:


get_ipython().system('pip install pyngrok')


# In[3]:


get_ipython().system('pip install translator')


# In[4]:


get_ipython().system('pip install googletrans')


# In[5]:


import os
import docx
from docx import Document
from translate import Translator
import streamlit as st


# In[7]:


from docx import Document

def load_demo_script(file_path):
    doc = Document(file_path)
    content = []
    for para in doc.paragraphs:
        content.append(para.text)
    return content

original_demo_script = load_demo_script("C:/Users/ADMIN/Downloads/demo-script-employee-central-overview.docx")


# In[8]:


import requests
import re
from googletrans import Translator

def get_google_translate_token():
    # Retrieve the Google Translate token using an HTTP request
    url = "https://translate.google.com"
    response = requests.get(url)
    tkk_expr = re.search("(tkk:.*?),", response.text)
    if tkk_expr:
        return tkk_expr.group(1)
    return None

def translate_content(content, target_language):
    # Get the Google Translate token
    token = get_google_translate_token()
    if not token:
        raise Exception("Failed to retrieve Google Translate token.")

    translator = Translator()
    translator.session.params["tk"] = token

    translated_content = [translator.translate(para, dest=target_language).text for para in content]
    return translated_content


# In[9]:


from googletrans import Translator

def translate_content(content, target_language):
    translator = Translator()
    translated_content = [translator.translate(para, dest=target_language).text for para in content]
    return translated_content

def replace_names(content, country):
    # Replace names based on the selected country
    # For example, replace "Alex Smith" with "Rahul Sharma" for India
    if country == "India":
        content = [para.replace("Alex Smith", "Rahul Sharma") for para in content]
    # Add more name replacements for other countries as needed
    return content

# Translate and replace names
localized_content = translate_content(original_demo_script, target_language="es")
localized_content = replace_names(localized_content, country="India")


# In[10]:


import streamlit as st

def main():
    st.title("Demo Script Localization")

    # File upload
    uploaded_file = st.file_uploader("Upload DOCX File", type=["docx"])

    if uploaded_file:
        target_language = st.selectbox("Select Target Language", ["Spanish", "Other Languages"])
        country = st.selectbox("Select Country", ["India", "Other Countries"])

        if st.button("Localize"):
            original_content = load_demo_script(uploaded_file)
            localized_content = translate_content(original_content, target_language)
            localized_content = replace_names(localized_content, country)

            # Save localized content to a new DOCX file
            localized_doc = Document()
            for para in localized_content:
                localized_doc.add_paragraph(para)
            localized_doc.save("Localized_Demo_Script.docx")

            st.success("Localization complete. Download the localized file.")
            st.download_button("Download Localized Script", "Localized_Demo_Script.docx")

if __name__ == "__main__":
    main()


# In[11]:


get_ipython().system('pip install python-docx translate streamlit')
get_ipython().system('pip install pyngrok')
get_ipython().system('pip install googletrans')

import re
import requests
from googletrans import Translator
from docx import Document
import streamlit as st

def load_demo_script(file_path):
    doc = Document(file_path)
    content = [para.text for para in doc.paragraphs]
    return content

def get_google_translate_token():
    # Retrieve the Google Translate token using an HTTP request
    url = "https://translate.google.com"
    response = requests.get(url)
    tkk_expr = re.search(r"tkk:'(.*?)'", response.text)
    if tkk_expr:
        return tkk_expr.group(1)
    return None

def translate_content(content, target_language):
    # Get the Google Translate token
    token = get_google_translate_token()
    if not token:
        raise Exception("Failed to retrieve Google Translate token.")

    translator = Translator()
    translator.session.params["tk"] = token

    translated_content = [translator.translate(para, dest=target_language).text for para in content]
    return translated_content

def replace_names(content, country):
    # Replace names based on the selected country
    # For example, replace "Alex Smith" with "Rahul Sharma" for India
    if country == "India":
        content = [para.replace("Alex Smith", "Rahul Sharma") for para in content]
    # Add more name replacements for other countries as needed
    return content

def main():
    st.title("Demo Script Localization")

    # File upload
    uploaded_file = st.file_uploader("Upload DOCX File", type=["docx"])

    if uploaded_file:
        target_language = st.selectbox("Select Target Language", ["es", "Other Languages"])
        country = st.selectbox("Select Country", ["India", "Other Countries"])

        if st.button("Localize"):
            original_content = load_demo_script(uploaded_file)
            localized_content = translate_content(original_content, target_language)
            localized_content = replace_names(localized_content, country)

            # Save localized content to a new DOCX file
            localized_doc = Document()
            for para in localized_content:
                localized_doc.add_paragraph(para)
            localized_doc.save("Localized_Demo_Script.docx")

            st.success("Localization complete. Download the localized file.")
            st.download_button("Download Localized Script", "Localized_Demo_Script.docx")

if __name__ == "__main__":
    main()


# In[ ]:





# In[ ]:





# In[ ]:




